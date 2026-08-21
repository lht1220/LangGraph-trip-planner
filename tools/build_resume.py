from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "output" / "李海涛_AI_Agent_应用开发简历.docx"

FONT_LATIN = "Arial"
FONT_CJK = "Microsoft YaHei"
NAVY = "14304A"
BLUE = "1F5A7A"
TEAL = "0C7C86"
INK = "17212B"
MUTED = "5B6573"
LIGHT = "F3F6F8"
WHITE = "FFFFFF"


def set_run_font(run, size=None, bold=None, color=INK, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_paragraph_padding(paragraph, before=60, after=60, left=100, right=100):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:right"), str(right))
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))


def add_hyperlink(paragraph, text, url, color=TEAL):
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "18")
    r_pr.append(sz)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, text, end):
        wrapper = OxmlElement("w:r")
        wrapper.append(el)
        paragraph._p.append(wrapper)


def add_custom_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "432")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "432")
    ind.set(qn("w:hanging"), "230")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "38")
    spacing.set(qn("w:line"), "260")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def set_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId")
    n_id.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(n_id)


def add_section_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=11.5, bold=True, color=BLUE)
    return p


def add_labeled_line(doc, label, text, after=2.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label)
    set_run_font(r, size=9.35, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=9.35, color=INK)
    return p


def add_project_header(doc, title, date):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run(title)
    set_run_font(r, size=10.6, bold=True, color=NAVY)
    r = p.add_run("\t" + date)
    set_run_font(r, size=9.1, bold=True, color=MUTED)
    return p


def add_meta(doc, role, stack, link=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2.8)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(role + "  |  ")
    set_run_font(r, size=8.9, bold=True, color=TEAL)
    r = p.add_run(stack)
    set_run_font(r, size=8.9, color=MUTED)
    if link:
        r = p.add_run("  |  ")
        set_run_font(r, size=8.9, color=MUTED)
        add_hyperlink(p, link[0], link[1])
    return p


def add_bullet(doc, num_id, text, bold_fragments=()):
    p = doc.add_paragraph()
    set_numbering(p, num_id)
    p.paragraph_format.left_indent = Inches(0.30)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.9)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = True
    remaining = text
    while remaining:
        next_match = None
        next_index = len(remaining)
        for frag in bold_fragments:
            idx = remaining.find(frag)
            if idx >= 0 and idx < next_index:
                next_match = frag
                next_index = idx
        if next_match is None:
            r = p.add_run(remaining)
            set_run_font(r, size=9.25, color=INK)
            break
        if next_index:
            r = p.add_run(remaining[:next_index])
            set_run_font(r, size=9.25, color=INK)
        r = p.add_run(next_match)
        set_run_font(r, size=9.25, bold=True, color=NAVY)
        remaining = remaining[next_index + len(next_match):]
    return p


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(9.35)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.08

    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_LATIN
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    h1.font.size = Pt(11.5)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(3.5)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_LATIN
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    h2.font.size = Pt(10.6)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(4.8)
    h2.paragraph_format.space_after = Pt(1.5)
    h2.paragraph_format.keep_with_next = True


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("李海涛  ·  AI Agent / 大模型应用开发")
    set_run_font(hr, size=8, color=MUTED)

    for footer in (section.first_page_footer, section.footer):
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fr = fp.add_run("第 ")
        set_run_font(fr, size=8, color=MUTED)
        add_field(fp, "PAGE")
        fr = fp.add_run(" / ")
        set_run_font(fr, size=8, color=MUTED)
        add_field(fp, "NUMPAGES")
        fr = fp.add_run(" 页")
        set_run_font(fr, size=8, color=MUTED)


def add_header(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run("李海涛")
    set_run_font(r, size=23, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3.8)
    r = p.add_run("AI Agent 应用开发工程师  /  大模型应用开发工程师  /  Python 后端工程师")
    set_run_font(r, size=10.6, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("杭州  |  133 5388 6509  |  2740566382@qq.com  |  ")
    set_run_font(r, size=9.1, color=MUTED)
    add_hyperlink(p, "github.com/lht1220", "https://github.com/lht1220")

    p = doc.add_paragraph()
    set_cell_shading(p, LIGHT)
    set_paragraph_padding(p, before=60, after=60, left=120, right=120)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run("核心定位  ")
    set_run_font(r, size=9.35, bold=True, color=TEAL)
    r = p.add_run(
        "计算机硕士在读，具备 1.5 年企业项目开发经验。能够将业务建模、状态流转、第三方系统集成等后端工程能力，"
        "应用于 LangGraph 多智能体编排、MCP Tool Calling 与 AI 应用服务化落地。"
    )
    set_run_font(r, size=9.35, color=INK)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    configure_page(doc)
    num_id = add_custom_bullet_numbering(doc)
    add_header(doc)

    add_section_heading(doc, "教育背景")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("浙江理工大学  ·  计算机技术（硕士）")
    set_run_font(r, size=9.5, bold=True, color=NAVY)
    r = p.add_run("\t2024.09 - 至今")
    set_run_font(r, size=9.1, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.2)
    r = p.add_run("GPA 3.6/5，专业前 5%  |  机器学习、人工智能、程序设计理论、图像理解与分析")
    set_run_font(r, size=9.05, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run("郑州航空工业管理学院  ·  软件工程（本科）")
    set_run_font(r, size=9.5, bold=True, color=NAVY)
    r = p.add_run("\t2020.09 - 2024.06")
    set_run_font(r, size=9.1, bold=True, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.2)
    r = p.add_run("GPA 3.3/4，专业前 10%  |  数据结构与算法、Java Web、数据库、计算机网络")
    set_run_font(r, size=9.05, color=MUTED)

    add_section_heading(doc, "专业技能")
    add_labeled_line(doc, "Agent / LLM：", "LangGraph、LangChain、StateGraph、Typed State、Reducer、Multi-Agent、Tool Calling、MCP、Prompt Engineering、OpenAI Compatible API")
    add_labeled_line(doc, "Python / 后端：", "Python、FastAPI、Pydantic、asyncio；Java、Spring Boot、MyBatis、MySQL、Redis、RESTful API")
    add_labeled_line(doc, "AI 应用理解：", "了解 RAG、Embedding、Vector Database、Structured Output、Agent Memory、Human-in-the-loop；关注 Grounding 与模型幻觉治理")
    add_labeled_line(doc, "工程 / 前端：", "Linux、Git、第三方服务集成与部署排障；Vue3、TypeScript、Vite、Axios、Ant Design Vue；具备终端设备接口联调经验")

    add_section_heading(doc, "项目经历")
    add_project_header(doc, "基于 LangGraph 的多智能体智能旅行规划系统", "2026.03 - 至今")
    add_meta(
        doc,
        "个人项目 / AI Agent 全栈开发",
        "Python · LangGraph · LangChain · MCP · FastAPI · Pydantic · Vue3 · TypeScript",
        ("GitHub", "https://github.com/lht1220/LangGraph-trip-planner"),
    )
    add_bullet(doc, num_id, "基于 LangGraph StateGraph 编排景点、天气、酒店与规划 4 个专业 Agent，将独立检索任务由串行重构为并行分支，经 Join 节点汇聚 Research Result 后统一交由 Planner Agent 决策。", ("LangGraph StateGraph", "并行分支", "Join 节点"))
    add_bullet(doc, num_id, "以 Typed State 管理请求、检索结果、规划结果与异常信息，通过 Reducer 处理多并行节点对共享状态的更新及错误传播，提升工作流状态一致性。", ("Typed State", "Reducer"))
    add_bullet(doc, num_id, "基于 MCP 与 langchain-mcp-adapters 接入高德地图 POI、天气等工具，使 Agent 按上下文自主 Tool Calling；兼容异步 MCP Tool 与同步 Agent 调用链，增强外部数据 Grounding。", ("MCP", "Tool Calling", "Grounding"))
    add_bullet(doc, num_id, "使用 FastAPI + Pydantic 服务化封装 Agent Workflow，并抽象 OpenAI Compatible 模型层，支持模型、Base URL 灵活切换及 LLM / Tool 实例复用。", ("FastAPI + Pydantic", "OpenAI Compatible"))
    add_bullet(doc, num_id, "采用 LLM Planning + Deterministic Algorithm 混合架构：以 Haversine 距离和最近邻策略优化访问顺序，并结合天气、预算、强度生成每日评分，避免将确定性计算完全交给模型。", ("LLM Planning + Deterministic Algorithm", "Haversine", "最近邻策略"))
    add_bullet(doc, num_id, "基于 Vue3 + TypeScript 实现地图 Marker、路线、天气、预算、评分与行程调整；支持轻松模式、雨天备选、路线优化和景点替换，形成从 Agent 编排到交互端的完整闭环。", ("Vue3 + TypeScript", "完整闭环"))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)

    add_section_heading(doc, "项目经历（续）")
    add_project_header(doc, "锐通网络穿透 SaaS 平台", "2025.11 - 至今")
    add_meta(doc, "后端 / 全栈开发", "Java · Spring Boot · MyBatis · MySQL · Redis · Vue · Linux · 支付宝 · 阿里云", ("nexoraweb.com.cn", "https://nexoraweb.com.cn/"))
    add_bullet(doc, num_id, "参与平台核心业务设计，围绕用户、会员、映射管理和服务器资源分配建立业务模型；根据服务配置管理端口及服务器资源，支撑映射资源的有效分配。", ("核心业务设计", "服务器资源分配"))
    add_bullet(doc, num_id, "集成支付宝支付，完成订单创建、异步回调、会员权益更新与状态一致性处理；基于阿里云接口实现 DNS 解析自动增删，打通自定义域名与映射服务。", ("支付宝支付", "状态一致性", "DNS 解析自动增删"))
    add_bullet(doc, num_id, "接入阿里云 OSS 并设计关键操作日志与业务追溯机制；独立完成部分运营后台、前后端联调和 Linux 部署排障，覆盖从业务设计到上线维护的完整链路。", ("业务追溯", "Linux 部署排障"))

    add_project_header(doc, "浙江大学生物实验中心智慧实验室管理平台", "2025.04 - 2025.08")
    add_meta(doc, "Java 后端开发", "Java · Spring Boot · Spring Security · MyBatis · MySQL · Redis · RFID")
    add_bullet(doc, num_id, "负责试剂、耗材、申领流转与审核模块的数据模型和业务流程设计，围绕入库、申领、领用、归还构建多角色状态流，保障业务数据一致性与可追溯性。", ("数据模型", "多角色状态流", "一致性与可追溯性"))
    add_bullet(doc, num_id, "面向 Android 一体机设计试剂领用、归还、入库等接口，与 RFID / 扫码现场流程联动；参与 Spring Security 权限控制及多端异常联调。", ("RFID / 扫码", "Spring Security"))
    add_bullet(doc, num_id, "通过标准化接口和业务状态管理实现 Web 管理端、后台服务与实验室终端协同，为试剂全生命周期统计、审计和安全管控提供数据基础。", ("全生命周期", "审计"))

    add_section_heading(doc, "实习经历")
    add_project_header(doc, "杭州聚数智创科技有限公司  ·  后端开发工程师", "2024.07 - 2026.01")
    add_bullet(doc, num_id, "深度参与实验室管理、设备管理、考试系统等企业项目，从需求分析、领域对象与数据库设计、RESTful API 到上线维护，主要负责 Java 后端与核心业务流程。", ("领域对象与数据库设计", "核心业务流程"))
    add_bullet(doc, num_id, "参与 Web、Android 一体机、第三方硬件与后台服务的接口设计和联调，处理跨终端状态一致性、异常场景和历史兼容问题。", ("跨终端状态一致性",))
    add_bullet(doc, num_id, "集成支付宝、人脸识别、打印设备、RFID / 扫码等外部能力，形成依据技术文档快速完成第三方服务与硬件接入的工程经验。", ("第三方服务与硬件接入",))
    add_bullet(doc, num_id, "将业务建模、状态流转、系统集成和部署排障经验迁移到 AI Agent 方向，重点实践 Workflow、Tool Calling 及 AI 应用服务化。", ("AI Agent", "Workflow", "Tool Calling"))

    add_section_heading(doc, "证书与补充")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("大学英语四级（CET-4）  |  蓝桥杯获奖  |  熟悉 AI 辅助开发工作流，可用于需求拆解、方案设计、代码生成、调试重构与快速原型验证")
    set_run_font(r, size=9.15, color=INK)

    doc.core_properties.title = "李海涛 - AI Agent 应用开发简历"
    doc.core_properties.subject = "AI Agent / 大模型应用开发 / Python 后端"
    doc.core_properties.author = "李海涛"
    doc.core_properties.keywords = "LangGraph, LangChain, MCP, Agent, Python, FastAPI, Java"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
